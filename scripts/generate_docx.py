"""Генератор Word-документа со всеми возможностями Мебельного портала."""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_para(text):
    doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    doc.add_paragraph()

# ============================================================
# ТИТУЛЬНАЯ СТРАНИЦА
# ============================================================
add_title('МЕБЕЛЬНЫЙ ПОРТАЛ\nПолное описание возможностей и функций')
doc.add_paragraph()
add_para('Версия документа: 2.0')
add_para('Дата: 24 июня 2026 г.')
add_para('Стек: Next.js 14 + React 18 + TypeScript + Prisma + PostgreSQL (Neon) + Tailwind CSS + JWT')
add_para('Статус: MVP — стабилизация')
doc.add_page_break()

# ============================================================
# СОДЕРЖАНИЕ
# ============================================================
add_title('СОДЕРЖАНИЕ')
toc_items = [
    '1. Обзор сервиса',
    '2. Авторизация и профиль',
    '3. Главная страница (Homepage)',
    '4. Каталог изображений (Gallery)',
    '5. Каталог документов (Documents)',
    '6. Справочники (References)',
    '7. Лента новостей (Feed)',
    '8. Портфолио',
    '9. Профили участников',
    '10. Личный кабинет (Dashboard)',
    '11. Система подписок',
    '12. Уведомления',
    '13. Избранное и закладки',
    '14. Личные сообщения',
    '15. Группы и сообщества',
    '16. Мероприятия',
    '17. Каталог товаров',
    '18. Сравнение товаров',
    '19. Глобальный поиск',
    '20. Горячие клавиши',
    '21. Админ-панель',
    '22. API эндпоинты',
    '23. Prisma модели (21 модель)',
    '24. Компоненты UI (25 шт)',
    '25. Безопасность',
    '26. Текущие проблемы и планы',
]
for item in toc_items:
    add_para(item)
doc.add_page_break()

# ============================================================
# 1. ОБЗОР СЕРВИСА
# ============================================================
add_h1('1. Обзор сервиса')
add_para('Мебельный портал — комплексная online-платформа для мебельной отрасли, объединяющая поставщиков, производителей, компании и специалистов. Сервис предоставляет каталоги изображений и документов, ленту новостей, портфолио работ, личные сообщения, систему подписок, уведомлений, группы, события и каталог товаров.')

add_h2('Ключевые метрики')
add_table(
    ['Параметр', 'Значение'],
    [
        ['Страниц (frontend)', '46'],
        ['API-эндпоинтов', '58'],
        ['UI-компонентов', '25'],
        ['Prisma-моделей', '21'],
        ['Custom hooks', '1 (useHotkeys)'],
        ['Стек', 'Next.js 14 App Router + TypeScript + Tailwind CSS'],
        ['База данных', 'PostgreSQL (Neon) через Prisma ORM'],
        ['Авторизация', 'JWT-токены (cookie + localStorage) + jose (middleware)'],
        ['Rate Limiting', 'In-memory (login: 10/мин, register: 5/мин, upload: 20/мин)'],
    ]
)

add_h2('Целевая аудитория')
add_bullet('Поставщики мебельных материалов и фурнитуры')
add_bullet('Производители мебели')
add_bullet('Мебельные компании и фабрики')
add_bullet('Дизайнеры интерьеров')
add_bullet('Технологи мебельного производства')
add_bullet('Монтажники и установщики')
add_bullet('Менеджеры по продажам мебели')

add_h2('Уникальность')
add_bullet('Единственная российская платформа, объединяющая всю мебельную отрасль')
add_bullet('Бесплатный доступ к библиотеке изображений, документам и справочникам')
add_bullet('Социальные функции: лента, группы, сообщения, портфолио')
add_bullet('Каталог товаров с отзывами и сравнением')

# ============================================================
# 2. АВТОРИЗАЦИЯ И ПРОФИЛЬ
# ============================================================
add_h1('2. Авторизация и профиль')

add_h2('2.1 Регистрация')
add_bullet('Email + пароль (минимум 6 символов, Zod-валидация)')
add_bullet('Хэширование пароля (bcryptjs, 12 раундов)')
add_bullet('JWT-токен создаётся при регистрации (срок: 7 дней)')
add_bullet('Роли при регистрации: USER, COMPANY, SUPPLIER, MANUFACTURER')
add_bullet('Автоматическое создание связанной сущности (Company/Supplier/Manufacturer)')
add_bullet('Необязательные поля: имя, ИНН')
add_bullet('Rate limiting: 5 попыток/мин')

add_h2('2.2 Вход')
add_bullet('Email + пароль')
add_bullet('Rate limiting: 10 попыток/мин')
add_bullet('JWT токен сохраняется в cookie + localStorage')
add_bullet('Middleware проверяет JWT через jose (Edge Runtime совместимость)')
add_bullet('Перенаправление на redirect-страницу после входа')

add_h2('2.3 Профиль пользователя')
add_bullet('Имя, email, телефон, ИНН, аватар')
add_bullet('Роль: USER, SPECIALIST, COMPANY, SUPPLIER, MANUFACTURER, ADMIN')
add_bullet('Привязка к бизнес-сущности (company/supplier/manufacturer/specialist)')
add_bullet('Поле: specialistId, companyId, supplierId, manufacturerId')
add_bullet('Дата регистрации (createdAt, updatedAt)')

add_h2('2.4 API авторизации')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/auth/register', 'POST', 'Регистрация нового пользователя'],
        ['/api/auth/login', 'POST', 'Вход (возвращает user + token)'],
        ['/api/auth/me', 'GET', 'Текущий пользователь (по токену)'],
        ['/api/auth/update', 'PUT', 'Обновление профиля'],
    ]
)

# ============================================================
# 3. ГЛАВНАЯ СТРАНИЦА
# ============================================================
add_h1('3. Главная страница (Homepage)')

add_h2('3.1 Секции homepage')
add_bullet('Hero-секция с градиентным фоном и CTA-кнопками')
add_bullet('Динамическая статистика (AnimatedCounter с IntersectionObserver)')
add_bullet('Популярные изображения (сетка 6 колонок)')
add_bullet('Лента последних постов (4 карточки с фото, лайками, комментариями)')
add_bullet('Секция "Всё для мебельщика" (3 карточки: изображения, документы, поставщики)')
add_bullet('Популярные документы (список с иконками по типу файла)')
add_bullet('Отзывы пользователей (3 карточки с рейтингом)')
add_bullet('CTA-секция «Присоединяйтесь к профессионалам»')

add_h2('3.2 Динамические данные')
add_bullet('GET /api/stats — количество изображений, документов, компаний, пользователей')
add_bullet('GET /api/feed?limit=4&sort=newest — последние посты')
add_bullet('GET /api/images?limit=6&sort=downloads — популярные изображения')
add_bullet('GET /api/documents?limit=5 — популярные документы')

# ============================================================
# 4. КАТАЛОГ ИЗОБРАЖЕНИЙ
# ============================================================
add_h1('4. Каталог изображений (Gallery)')

add_h2('4.1 Список изображений')
add_bullet('Сетка с responsive breakpoints (2/3/4/6 колонок)')
add_bullet('Сортировка: newest, popular/downloads')
add_bullet('Фильтрация по стилям и категориям')
add_bullet('Постраничная навигация')
add_bullet('Скелетон-загрузка')

add_h2('4.2 Карточка изображения')
add_bullet('Превью изображения с hover-эффектом')
add_bullet('Название, стиль, категория')
add_bullet('Теги (badge-стиль)')
add_bullet('Количество скачиваний')
add_bullet('Кнопка «Скачать»')
add_bullet('Кнопка «Избранное» (FavoriteButton)')

add_h2('4.3 Детальная страница (/gallery/[id])')
add_bullet('Полное изображение')
add_bullet('Метаданные: описание, стиль, категория, теги')
add_bullet('Статистика скачиваний')
add_bullet('Похожие изображения')

add_h2('4.4 API изображений')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/images', 'GET', 'Список с фильтрами/сортировкой/пагинацией'],
        ['/api/images/[id]', 'GET', 'Детали изображения'],
        ['/api/images/upload', 'POST', 'Загрузка изображения'],
        ['/api/images/[id]', 'DELETE', 'Удаление изображения'],
    ]
)

# ============================================================
# 5. КАТАЛОГ ДОКУМЕНТОВ
# ============================================================
add_h1('5. Каталог документов (Documents)')

add_h2('5.1 Функционал')
add_bullet('Поиск по названию')
add_bullet('Сортировка: newest, popular/downloads')
add_bullet('Иконки по типу файла (PDF, DOC, XLS)')
add_bullet('Inline скачивание')
add_bullet('Категории: Договоры, Акты, Спецификации, ТЗ, Счета')
add_bullet('Постраничная навигация')

add_h2('5.2 API документов')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/documents', 'GET', 'Список с фильтрами/сортировкой'],
        ['/api/documents/[id]', 'GET', 'Детали документа'],
    ]
)

# ============================================================
# 6. СПРАВОЧНИКИ
# ============================================================
add_h1('6. Справочники (References)')

add_h2('6.1 Категории')
add_bullet('Размеры (стандартные размеры техники, кромкования)')
add_bullet('Фурнитура (паспорта фурнитуры Blum, Hettich)')
add_bullet('Нормы (расход материалов)')
add_bullet('Материалы (сравнение материалов фасадов)')

add_h2('6.2 Функционал')
add_bullet('Табличное отображение данных (RefTable компонент)')
add_bullet('Поиск по содержимому')
add_bullet('Категоризированные справочные материалы')
add_bullet('Интерактивные таблицы с.expand-строками')

add_h2('6.3 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/refs', 'GET', 'Список справочников с фильтрацией'],
    ]
)

# ============================================================
# 7. ЛЕНТА НОВОСТЕЙ
# ============================================================
add_h1('7. Лента новостей (Feed)')

add_h2('7.1 Просмотр ленты')
add_bullet('Карточки постов с фото, автором, тегами')
add_bullet('Категории: news, project, article, product')
add_bullet('Время создания (относительное: "5 мин. назад")')
add_bullet('Просмотры, лайки, комментарии')
add_bullet('Кнопка «Поделиться» (копирование ссылки)')

add_h2('7.2 Фильтрация')
add_bullet('По категории (news/project/article/product)')
add_bullet('По подпискам (filter=subscriptions)')
add_bullet('По автору (authorId)')
add_bullet('Сортировка: newest, popular, discussed')
add_bullet('Пагинация (page, limit)')

add_h2('7.3 Создание поста (/feed/new)')
add_bullet('Заголовок и содержание (до 5000 символов)')
add_bullet('Выбор категории')
add_bullet('Загрузка изображений (до 10 шт, через /api/upload)')
add_bullet('Добавление тегов (до 10, макс. 30 символов)')
add_bullet('Черновик / публикация (isPublished)')

add_h2('7.4 Детали поста (/feed/[id])')
add_bullet('Галерея изображений с лайтбоксом')
add_bullet('Автор с аватаром и ссылкой на профиль')
add_bullet('Кнопки: лайк, комментарий, редактирование (для автора)')
add_bullet('Список комментариев')
add_bullet('Форма добавления комментария')

add_h2('7.5 Редактирование (/feed/[id]/edit)')
add_bullet('Изменение заголовка, содержания, категории, тегов')
add_bullet('Добавление/удаление изображений')
add_bullet('Проверка авторства (только автор или ADMIN)')

add_h2('7.6 Лайки')
add_bullet('Toggle лайка (POST /api/posts/[id]/like)')
add_bullet('Уникальность: @@unique([userId, postId])')
add_bullet('Счётчик лайков на посте')

add_h2('7.7 Репосты')
add_bullet('Репост с комментарием (POST /api/posts/[id]/repost)')
add_bullet('Уникальность: @@unique([userId, postId])')

add_h2('7.8 Комментарии')
add_bullet('Добавление комментариев (POST /api/posts/[id]/comments)')
add_bullet('Список комментариев с аватарами авторов')
add_bullet('Каскадное удаление (onDelete: Cascade)')

# ============================================================
# 8. ПОРТФОЛИО
# ============================================================
add_h1('8. Портфолио')

add_h2('8.1 Личное портфолио (Dashboard)')
add_bullet('/dashboard/portfolio — список работ')
add_bullet('/dashboard/portfolio/new — добавление новой работы')
add_bullet('/dashboard/portfolio/[id]/edit — редактирование')
add_bullet('Загрузка изображений работ')
add_bullet('Категории и теги')
add_bullet('Черновик / публикация')

add_h2('8.2 Публичное портфолио')
add_bullet('/portfolio/[userId] — публичная страница портфолио')
add_bullet('Сетка работ с фильтрами по категориям')
add_bullet('Не требует авторизации')

add_h2('8.3 API портфолио')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/portfolio', 'GET', 'Портфолио текущего пользователя'],
        ['/api/portfolio', 'POST', 'Создание элемента портфолио'],
        ['/api/portfolio/[id]', 'PUT', 'Редактирование элемента'],
        ['/api/portfolio/[id]', 'DELETE', 'Удаление элемента'],
        ['/api/portfolio/user/[userId]', 'GET', 'Публичное портфолио пользователя'],
    ]
)

# ============================================================
# 9. ПРОФИЛИ УЧАСТНИКОВ
# ============================================================
add_h1('9. Профили участников')

add_h2('9.1 Специалисты (/specialists)')
add_bullet('Типы: менеджер, дизайнер, технолог, монтажник')
add_bullet('Карточки с рейтингом (звёзды), опытом')
add_bullet('Поиск по имени')
add_bullet('Сортировка: rating, experience, newest')
add_bullet('Detail-страница: описание, рейтинг, портфолио, посты')
add_bullet('Оценка специалиста (Rating, 1-5 звёзд)')

add_h2('9.2 Поставщики (/suppliers)')
add_bullet('Карточки с логотипом, категориями, верификацией')
add_bullet('Сортировка: newest, verified, products')
add_bullet('Detail-страница: контакты, продукция, посты')

add_h2('9.3 Компании (/companies)')
add_bullet('Карточки с логотипом, описанием, верификацией')
add_bullet('Сортировка: newest, verified, name')
add_bullet('Detail-страница: контакты, продукция, посты')
add_bullet('PUT /api/companies/[id] — редактирование (owner-only)')

add_h2('9.4 Производители (/manufacturers)')
add_bullet('Карточки с логотипом, возможностями, геометрией')
add_bullet('Detail-страница: контакты, возможности, посты')

add_h2('9.5 API участников')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/specialists', 'GET', 'Список специалистов'],
        ['/api/specialists/[id]', 'GET', 'Профиль специалиста'],
        ['/api/specialists/[id]/rate', 'POST', 'Оценка специалиста'],
        ['/api/suppliers', 'GET', 'Список поставщиков'],
        ['/api/suppliers/[id]', 'GET', 'Профиль поставщика'],
        ['/api/companies', 'GET', 'Список компаний'],
        ['/api/companies/[id]', 'GET', 'Профиль компании'],
        ['/api/companies/[id]', 'PUT', 'Редактирование компании'],
        ['/api/manufacturers', 'GET', 'Список производителей'],
        ['/api/manufacturers/[id]', 'GET', 'Профиль производителя'],
    ]
)

# ============================================================
# 10. ЛИЧНЫЙ КАБИНЕТ
# ============================================================
add_h1('10. Личный кабинет (Dashboard)')

add_h2('10.1 Основной экран (/dashboard)')
add_bullet('Gradient hero-секция с именем и статистикой')
add_bullet('Счётчики: скачивания, посты, портфолио, подписчики, подписки')
add_bullet('Быстрые ссылки: новый пост, добавить портфолио')

add_h2('10.2 Профиль (/dashboard/profile)')
add_bullet('Редактирование имени, телефона, ИНН')
add_bullet('Загрузка аватара')

add_h2('10.3 Настройки (/dashboard/settings)')
add_bullet('Настройки уведомлений')
add_bullet('Кнопка удаления аккаунта')

add_h2('10.4 Избранное (/dashboard/favorites)')
add_bullet('Все избранные элементы (изображения, документы, посты)')

add_h2('10.5 Закладки (/dashboard/bookmarks)')
add_bullet('Создание коллекций закладок')
add_bullet('Добавление элементов в закладки')
add_bullet('Управление коллекциями')

add_h2('10.6 Сообщения (/dashboard/messages)')
add_bullet('Список бесед с превью последнего сообщения')
add_bullet('Чат (/dashboard/messages/[id])')
add_bullet('Статус прочтения')

add_h2('10.7 Портфолио (/dashboard/portfolio)')
add_bullet('Управление элементами портфолио')

add_h2('10.8 Аналитика (/dashboard/analytics)')
add_bullet('Статистика активности пользователя')

# ============================================================
# 11. ПОДПИСКИ
# ============================================================
add_h1('11. Система подписок')

add_h2('11.1 Функционал')
add_bullet('Подписка / отписка (toggle)')
add_bullet('Счётчики подписчиков и подписок')
add_bullet('Фильтр ленты «Подписки»')
add_bullet('FollowButton: compact (для постов) + full (для профилей)')

add_h2('11.2 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/users/[id]/follow', 'POST', 'Toggle подписки'],
        ['/api/users/[id]/followers', 'GET', 'Список подписчиков'],
        ['/api/users/[id]/following', 'GET', 'Список подписок'],
        ['/api/users/[id]/follow-status', 'GET', 'Статус подписки'],
    ]
)

# ============================================================
# 12. УВЕДОМЛЕНИЯ
# ============================================================
add_h1('12. Уведомления')

add_h2('12.1 Типы уведомлений')
add_bullet('like — кто-то лайкнул пост')
add_bullet('comment — новый комментарий')
add_bullet('follow — новая подписка')
add_bullet('message — новое сообщение')

add_h2('12.2 Функционал')
add_bullet('Dropdown в Header (NotificationsDropdown)')
add_bullet('Счётчик непрочитанных (badge)')
add_bullet('Автоматическое создание при лайке/комментарии/подписке')
add_bullet('Пометка как прочитанного')
add_bullet('Poll каждые 30 секунд')

add_h2('12.3 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/notifications', 'GET', 'Список уведомлений'],
        ['/api/notifications/unread', 'GET', 'Количество непрочитанных'],
        ['/api/notifications/[id]/read', 'PUT', 'Пометить как прочитанное'],
    ]
)

# ============================================================
# 13. ИЗБРАННОЕ И ЗАКЛАДКИ
# ============================================================
add_h1('13. Избранное и закладки')

add_h2('13.1 Избранное')
add_bullet('Добавление / удаление из избранного (toggle)')
add_bullet('Полиморфная модель: itemType (image/document/post/portfolio)')
add_bullet('Уникальность: @@unique([userId, itemType, itemId])')
add_bullet('FavoriteButton на карточках и detail-страницах')
add_bullet('Страница /dashboard/favorites')

add_h2('13.2 Закладки')
add_bullet('Создание именованных коллекций (Bookmark)')
add_bullet('Добавление элементов в коллекции (BookmarkItem)')
add_bullet('Полиморфная модель: itemType + itemId')
add_bullet('Уникальность: @@unique([bookmarkId, itemType, itemId])')
add_bullet('Страница /dashboard/bookmarks')

add_h2('13.3 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/favorites', 'GET', 'Список избранного'],
        ['/api/favorites', 'POST', 'Toggle избранного'],
        ['/api/bookmarks', 'GET', 'Список коллекций'],
        ['/api/bookmarks', 'POST', 'Создать коллекцию'],
        ['/api/bookmarks/[id]', 'GET', 'Коллекция с элементами'],
        ['/api/bookmarks/[id]', 'DELETE', 'Удалить коллекцию'],
        ['/api/bookmarks/[id]/items', 'POST', 'Добавить элемент'],
    ]
)

# ============================================================
# 14. ЛИЧНЫЕ СООБЩЕНИЯ
# ============================================================
add_h1('14. Личные сообщения')

add_h2('14.1 Функционал')
add_bullet('Беседы между пользователями')
add_bullet('Список бесед (/dashboard/messages)')
add_bullet('Чат (/dashboard/messages/[id])')
add_bullet('Отправка сообщений')
add_bullet('Статус прочтения (lastReadAt)')
add_bullet('Кнопка «Написать сообщение» на профилях')

add_h2('14.2 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/conversations', 'GET', 'Список бесед'],
        ['/api/conversations', 'POST', 'Создание беседы'],
        ['/api/conversations/[id]/messages', 'GET', 'Сообщения беседы'],
        ['/api/conversations/[id]/messages', 'POST', 'Отправка сообщения'],
    ]
)

# ============================================================
# 15. ГРУППЫ И СООБЩЕСТВА
# ============================================================
add_h1('15. Группы и сообщества')

add_h2('15.1 Функционал')
add_bullet('Создание групп (публичные/приватные)')
add_bullet('Обложка, описание')
add_bullet('Owner + участники с ролями')
add_bullet('Посты в группах (GroupPost)')
add_bullet('Вступление/выход из группы')
add_bullet('Список участников')

add_h2('15.2 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/groups', 'GET', 'Список групп'],
        ['/api/groups', 'POST', 'Создание группы'],
        ['/api/groups/[id]', 'GET', 'Детали группы'],
        ['/api/groups/[id]', 'PUT', 'Редактирование группы'],
        ['/api/groups/[id]/join', 'POST', 'Вступить/выйти из группы'],
        ['/api/groups/[id]/posts', 'GET', 'Посты группы'],
        ['/api/groups/[id]/posts', 'POST', 'Создать пост в группе'],
    ]
)

# ============================================================
# 16. МЕРОПРИЯТИЯ
# ============================================================
add_h1('16. Мероприятия')

add_h2('16.1 Функционал')
add_bullet('Создание мероприятий')
add_bullet('Типы: offline, online, webinar')
add_bullet('Дата начала/окончания')
add_bullet('Локация')
add_bullet('Ограничение участников (maxParticipants)')
add_bullet('Запись на мероприятие')
add_bullet('Статус участия: going, maybe, interested')

add_h2('16.2 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/events', 'GET', 'Список мероприятий'],
        ['/api/events', 'POST', 'Создание мероприятия'],
        ['/api/events/[id]', 'GET', 'Детали мероприятия'],
        ['/api/events/[id]', 'PUT', 'Редактирование'],
        ['/api/events/[id]/join', 'POST', 'Запись на мероприятие'],
    ]
)

# ============================================================
# 17. КАТАЛОГ ТОВАРОВ
# ============================================================
add_h1('17. Каталог товаров')

add_h2('17.1 Функционал')
add_bullet('Каталог товаров с ценами и характеристиками')
add_bullet('Категории: Фурнитура, ЛДСП, МДФ, Техника')
add_bullet('Привязка к компании/поставщику')
add_bullet('Отзывы с рейтингом (1-5 звёзд)')
add_bullet('Сравнение товаров (/products/compare)')

add_h2('17.2 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/products', 'GET', 'Список товаров'],
        ['/api/products/[id]', 'GET', 'Детали товара'],
        ['/api/products/[id]/reviews', 'GET', 'Отзывы на товар'],
        ['/api/products/[id]/reviews', 'POST', 'Добавить отзыв'],
    ]
)

# ============================================================
# 18. СРАВНЕНИЕ ТОВАРОВ
# ============================================================
add_h1('18. Сравнение товаров')

add_h2('18.1 Функционал')
add_bullet('Страница /products/compare')
add_bullet('Выбор товаров для сравнения')
add_bullet('Таблица сравнения характеристик')
add_bullet('Цены, описания, характеристики бок о бок')

# ============================================================
# 19. ГЛОБАЛЬНЫЙ ПОИСК
# ============================================================
add_h1('19. Глобальный поиск')

add_h2('19.1 Функционал')
add_bullet('Ctrl+K — открытие модального окна поиска')
add_bullet('Поиск по: постам, изображениям, документам, товарам, компаниям, поставщикам, производителям, специалистам')
add_bullet('Мгновенные результаты')
add_bullet('Ссылки на найденные элементы')
add_bullet('Компонент SearchModal')

add_h2('19.2 API')
add_table(
    ['Эндпоинт', 'Метод', 'Описание'],
    [
        ['/api/search?q=...', 'GET', 'Глобальный поиск по всем сущностям'],
    ]
)

# ============================================================
# 20. ГОРЯЧИЕ КЛАВИШИ
# ============================================================
add_h1('20. Горячие клавиши')

add_h2('20.1 Функционал')
add_bullet('Ctrl+K — открытие поиска')
add_bullet('Ctrl+/ — список горячих клавиш')
add_bullet('Навигация: G then H (home), G then G (gallery), G then F (feed)')
add_bullet('Страница /shortcuts — справочник всех клавиш')
add_bullet('HotkeysProvider — глобальный провайдер')

add_h2('20.2 Реализация')
add_bullet('Custom hook: useHotkeys')
add_bullet('HotkeysProvider для глобального контекста')
add_bullet('Underlay с подсказками')

# ============================================================
# 21. АДМИН-ПАНЕЛЬ
# ============================================================
add_h1('21. Админ-панель')

add_h2('21.1 Дашборд (/admin)')
add_bullet('Общая статистика: пользователи, изображения, документы, посты')
add_bullet('Графики регистраций за 30 дней')
add_bullet('Последние регистрации')
add_bullet('Популярный контент')

add_h2('21.2 Управление пользователями (/admin/users)')
add_bullet('Список всех пользователей')
add_bullet('Роль, дата регистрации')
add_bullet('Изменение ролей')
add_bullet('Удаление')

add_h2('21.3 Управление изображениями (/admin/images)')
add_bullet('Загрузка через UI')
add_bullet('Удаление изображений')
add_bullet('Статистика скачиваний')

add_h2('21.4 Управление документами (/admin/documents)')
add_bullet('Список документов')
add_bullet('Удаление')

add_h2('21.5 Управление постами (/admin/posts)')
add_bullet('Список всех постов')
add_bullet('Модерация (isPublished toggle)')
add_bullet('Удаление постов')

add_h2('21.6 Управление компаниями (/admin/companies)')
add_bullet('Список компаний')
add_bullet('Верификация')

add_h2('21.7 Middleware защита')
add_bullet('/dashboard — требует авторизацию (JWT валидация через jose)')
add_bullet('/admin — требует роль ADMIN')
add_bullet('Невалидный токен → редирект на /login + удаление cookie')

# ============================================================
# 22. API ЭНДПОИНТЫ
# ============================================================
add_h1('22. Полный список API эндпоинтов')

api_groups = [
    ('Авторизация', [
        ['/api/auth/register', 'POST', 'Регистрация'],
        ['/api/auth/login', 'POST', 'Вход'],
        ['/api/auth/me', 'GET', 'Текущий пользователь'],
        ['/api/auth/update', 'PUT', 'Обновление профиля'],
    ]),
    ('Посты (Feed)', [
        ['/api/feed', 'GET', 'Лента постов'],
        ['/api/posts', 'GET', 'Список постов'],
        ['/api/posts', 'POST', 'Создание поста'],
        ['/api/posts/[id]', 'GET', 'Детали поста'],
        ['/api/posts/[id]', 'PUT', 'Редактирование'],
        ['/api/posts/[id]', 'DELETE', 'Удаление'],
        ['/api/posts/[id]/like', 'POST', 'Лайк/дизлайк'],
        ['/api/posts/[id]/repost', 'POST', 'Репост'],
        ['/api/posts/[id]/comments', 'GET', 'Комментарии'],
        ['/api/posts/[id]/comments', 'POST', 'Добавить комментарий'],
    ]),
    ('Портфолио', [
        ['/api/portfolio', 'GET', 'Моё портфолио'],
        ['/api/portfolio', 'POST', 'Создать элемент'],
        ['/api/portfolio/[id]', 'PUT', 'Редактировать'],
        ['/api/portfolio/[id]', 'DELETE', 'Удалить'],
        ['/api/portfolio/user/[userId]', 'GET', 'Публичное портфолио'],
    ]),
    ('Пользователи', [
        ['/api/users', 'GET', 'Список пользователей'],
        ['/api/users/[id]/follow', 'POST', 'Toggle подписки'],
        ['/api/users/[id]/followers', 'GET', 'Подписчики'],
        ['/api/users/[id]/following', 'GET', 'Подписки'],
        ['/api/users/[id]/follow-status', 'GET', 'Статус подписки'],
    ]),
    ('Организации', [
        ['/api/companies', 'GET', 'Список компаний'],
        ['/api/companies/[id]', 'GET', 'Профиль'],
        ['/api/companies/[id]', 'PUT', 'Редактирование'],
        ['/api/suppliers', 'GET', 'Список поставщиков'],
        ['/api/suppliers/[id]', 'GET', 'Профиль'],
        ['/api/manufacturers', 'GET', 'Список производителей'],
        ['/api/manufacturers/[id]', 'GET', 'Профиль'],
        ['/api/specialists', 'GET', 'Список специалистов'],
        ['/api/specialists/[id]', 'GET', 'Профиль'],
        ['/api/specialists/[id]/rate', 'POST', 'Оценка'],
    ]),
    ('Каталог', [
        ['/api/images', 'GET', 'Список изображений'],
        ['/api/images/[id]', 'GET', 'Детали'],
        ['/api/images/upload', 'POST', 'Загрузка'],
        ['/api/documents', 'GET', 'Список документов'],
        ['/api/documents/[id]', 'GET', 'Детали'],
        ['/api/refs', 'GET', 'Справочники'],
    ]),
    ('Товары', [
        ['/api/products', 'GET', 'Список товаров'],
        ['/api/products/[id]', 'GET', 'Детали товара'],
        ['/api/products/[id]/reviews', 'GET', 'Отзывы'],
        ['/api/products/[id]/reviews', 'POST', 'Добавить отзыв'],
    ]),
    ('Социальное', [
        ['/api/favorites', 'GET', 'Избранное'],
        ['/api/favorites', 'POST', 'Toggle'],
        ['/api/bookmarks', 'GET', 'Коллекции'],
        ['/api/bookmarks', 'POST', 'Создать'],
        ['/api/bookmarks/[id]', 'GET', 'Элементы'],
        ['/api/bookmarks/[id]', 'DELETE', 'Удалить'],
        ['/api/bookmarks/[id]/items', 'POST', 'Добавить элемент'],
        ['/api/notifications', 'GET', 'Уведомления'],
        ['/api/notifications/unread', 'GET', 'Непрочитанные'],
        ['/api/notifications/[id]/read', 'PUT', 'Прочитано'],
        ['/api/conversations', 'GET', 'Беседы'],
        ['/api/conversations', 'POST', 'Создать беседу'],
        ['/api/conversations/[id]/messages', 'GET', 'Сообщения'],
        ['/api/conversations/[id]/messages', 'POST', 'Отправить'],
    ]),
    ('Группы', [
        ['/api/groups', 'GET', 'Список групп'],
        ['/api/groups', 'POST', 'Создать группу'],
        ['/api/groups/[id]', 'GET', 'Детали'],
        ['/api/groups/[id]', 'PUT', 'Редактирование'],
        ['/api/groups/[id]/join', 'POST', 'Вступить/выйти'],
        ['/api/groups/[id]/posts', 'GET', 'Посты группы'],
        ['/api/groups/[id]/posts', 'POST', 'Создать пост'],
    ]),
    ('Мероприятия', [
        ['/api/events', 'GET', 'Список'],
        ['/api/events', 'POST', 'Создать'],
        ['/api/events/[id]', 'GET', 'Детали'],
        ['/api/events/[id]', 'PUT', 'Редактирование'],
        ['/api/events/[id]/join', 'POST', 'Записаться'],
    ]),
    ('Системные', [
        ['/api/stats', 'GET', 'Статистика (public)'],
        ['/api/upload', 'POST', 'Загрузка файлов'],
        ['/api/downloads', 'GET', 'История скачиваний'],
        ['/api/search', 'GET', 'Глобальный поиск'],
        ['/api/feed', 'GET', 'Лента'],
    ]),
    ('Админ', [
        ['/api/admin/stats', 'GET', 'Админ-статистика'],
        ['/api/admin/posts', 'GET', 'Управление постами'],
        ['/api/admin/posts', 'PATCH', 'Модерация поста'],
        ['/api/admin/posts', 'DELETE', 'Удаление поста'],
    ]),
]

for group_name, endpoints in api_groups:
    add_h2(group_name)
    add_table(['Эндпоинт', 'Метод', 'Описание'], endpoints)

# ============================================================
# 23. PRISMA МОДЕЛИ (21)
# ============================================================
add_h1('23. Prisma модели (База данных — 21 модель)')

models = [
    ('User', 'Пользователь системы', [
        'id (String, PK)', 'email (String, unique)', 'password (String, bcrypt)',
        'name (String?)', 'avatar (String?)', 'role (String, default USER)',
        'inn (String?)', 'phone (String?)', 'createdAt', 'updatedAt',
        'companyId?', 'specialistId? (unique)', 'supplierId?', 'manufacturerId?',
        'Связи: downloads, ratings, posts, portfolio, postLikes, comments,',
        'following, followers, notifications, favorites, conversations,',
        'sentMessages, ownedGroups, groupMemberships, groupPosts, bookmarks,',
        'reposts, organizedEvents, eventParticipations, productReviews'
    ]),
    ('Specialist', 'Специалист', [
        'id (String, PK)', 'type (String)', 'description (String?)',
        'experience (Int?)', 'portfolio (String?)', 'rating (Float, default 0)',
        'Связи: user (1:1), ratings'
    ]),
    ('Company', 'Компания', [
        'id (String, PK)', 'name (String)', 'description (String?)',
        'logo (String?)', 'website (String?)', 'phone (String?)',
        'email (String?)', 'address (String?)', 'isVerified (Boolean)',
        'Связи: users, products, posts'
    ]),
    ('Supplier', 'Поставщик', [
        'id (String, PK)', 'companyName (String)', 'description (String?)',
        'logo (String?)', 'website (String?)', 'phone (String?)',
        'email (String?)', 'categories (String, JSON)', 'isVerified (Boolean)',
        'Связи: products, users, posts'
    ]),
    ('Manufacturer', 'Производитель', [
        'id (String, PK)', 'name (String)', 'description (String?)',
        'logo (String?)', 'website (String?)', 'phone (String?)',
        'email (String?)', 'address (String?)', 'capabilities (String, JSON)',
        'geometry (String?)', 'isVerified (Boolean)',
        'Связи: users'
    ]),
    ('Product', 'Товар', [
        'id (String, PK)', 'name (String)', 'description (String?)',
        'price (Float?)', 'images (String, JSON)', 'category (String)',
        'brand (String?)', 'specs (String, JSON)', 'isPublished (Boolean)',
        'companyId?', 'supplierId?',
        'Связи: company, supplier, reviews'
    ]),
    ('ProductReview', 'Отзыв на товар', [
        'id (String, PK)', 'score (Int)', 'comment (String?)',
        'userId', 'productId',
        'Уникальность: @@unique([userId, productId])',
        'Индекс: @@index([productId])'
    ]),
    ('Image', 'Изображение в каталоге', [
        'id (String, PK)', 'title (String)', 'url (String)',
        'thumbnail (String?)', 'style (String?)', 'category (String?)',
        'tags (String, JSON)', 'downloads (Int, default 0)',
        'Индекс: @@index([downloads])',
        'Связи: downloadsList'
    ]),
    ('Download', 'Скачивание', [
        'id (String, PK)', 'userId', 'imageId',
        'Индексы: @@index([userId]), @@index([imageId])'
    ]),
    ('Document', 'Документ', [
        'id (String, PK)', 'title (String)', 'fileUrl (String)',
        'category (String)', 'fileType (String)', 'downloads (Int)'
    ]),
    ('Reference', 'Справочник', [
        'id (String, PK)', 'title (String)', 'content (String)', 'category (String)'
    ]),
    ('Post', 'Пост ленты', [
        'id (String, PK)', 'title (String)', 'content (String)',
        'category (String, default news)', 'images (String, JSON)',
        'tags (String, JSON)', 'likes (Int)', 'views (Int)',
        'isPublished (Boolean)', 'authorId', 'companyId?', 'supplierId?',
        'Индексы: @@index([authorId]), @@index([isPublished, createdAt])',
        'Связи: author, company, supplier, comments, likesList, reposts'
    ]),
    ('PortfolioItem', 'Элемент портфолио', [
        'id (String, PK)', 'title (String)', 'description (String?)',
        'imageUrl (String?)', 'images (String, JSON)', 'category (String?)',
        'tags (String, JSON)', 'isPublished (Boolean)', 'userId',
        'Индекс: @@index([userId])'
    ]),
    ('Comment', 'Комментарий', [
        'id (String, PK)', 'content (String)', 'authorId', 'postId',
        'onDelete: Cascade, @@index([postId])'
    ]),
    ('PostLike', 'Лайк поста', [
        'id (String, PK)', 'userId', 'postId',
        '@@unique([userId, postId]), @@index([postId])'
    ]),
    ('Follow', 'Подписка', [
        'id (String, PK)', 'followerId', 'followingId',
        '@@unique([followerId, followingId])'
    ]),
    ('Rating', 'Рейтинг специалиста', [
        'id (String, PK)', 'score (Int)', 'comment (String?)',
        'userId', 'specialistId',
        '@@unique([userId, specialistId]), @@index([specialistId])'
    ]),
    ('Favorite', 'Избранное', [
        'id (String, PK)', 'itemType (String)', 'itemId (String)', 'userId',
        '@@unique([userId, itemType, itemId])'
    ]),
    ('Conversation + ConversationParticipant + Message', 'Мессенджер', [
        'Conversation: id, createdAt, updatedAt',
        'Participant: userId, conversationId, lastReadAt',
        '@@unique([userId, conversationId])',
        'Message: content, authorId, conversationId',
        'Индексы: @@index([conversationId]), @@index([authorId]), @@index([createdAt])'
    ]),
    ('Group + GroupMember + GroupPost', 'Группы', [
        'Group: name, description, coverImage, type (public/private), ownerId',
        'GroupMember: userId, groupId, role (member)',
        '@@unique([userId, groupId]), onDelete: Cascade',
        'GroupPost: content, images, authorId, groupId',
        '@@index([groupId]), onDelete: Cascade'
    ]),
    ('Bookmark + BookmarkItem', 'Закладки', [
        'Bookmark: name, userId',
        'BookmarkItem: itemType, itemId, bookmarkId',
        '@@unique([bookmarkId, itemType, itemId]), onDelete: Cascade'
    ]),
    ('Repost', 'Репост', [
        'id (String, PK)', 'comment (String?)', 'userId', 'postId',
        '@@unique([userId, postId]), onDelete: Cascade'
    ]),
    ('Event + EventParticipant', 'Мероприятия', [
        'Event: title, description, coverImage, location, startDate, endDate,',
        'type (offline/online/webinar), maxParticipants, organizerId',
        'EventParticipant: status (going/maybe/interested), userId, eventId',
        '@@unique([userId, eventId]), @@index([eventId]), onDelete: Cascade'
    ]),
    ('Notification', 'Уведомление', [
        'id (String, PK)', 'type (String)', 'message (String)',
        'read (Boolean)', 'link (String?)', 'userId', 'fromUserId?', 'postId?',
        'Индексы: @@index([userId]), @@index([userId, read])'
    ]),
]

for model_name, description, fields in models:
    add_h2(f'{model_name} — {description}')
    for field in fields:
        add_bullet(field)

# ============================================================
# 24. КОМПОНЕНТЫ UI (25)
# ============================================================
add_h1('24. Компоненты UI (25 шт)')

components = [
    ('Header', 'Шапка: навигация (10 ссылок), уведомления, поиск, аватар, dropdown'),
    ('Footer', 'Подвал: ссылки на разделы, копирайт'),
    ('Sidebar', 'Боковая панель'),
    ('AuthForm', 'Форма входа/регистрации'),
    ('SearchModal', 'Модальное окно поиска (Ctrl+K)'),
    ('SearchBar', 'Строка поиска'),
    ('Pagination', 'Постраничная навигация'),
    ('Modal', 'Базовое модальное окно'),
    ('Toast', 'Всплывающие уведомления'),
    ('Loading', 'Skeleton загрузки'),
    ('FilterPanel', 'Панель фильтрации: категории, сортировка'),
    ('NotificationsDropdown', 'Выпадающий список уведомлений'),
    ('CompanyCard', 'Карточка компании'),
    ('SupplierCard', 'Карточка поставщика'),
    ('SpecialistCard', 'Карточка специалиста'),
    ('DocumentCard', 'Карточка документа'),
    ('ImageCard', 'Карточка изображения'),
    ('ImageGrid', 'Сетка изображений'),
    ('RefTable', 'Таблица справочника'),
    ('FollowButton', 'Кнопка подписки (compact + full)'),
    ('FavoriteButton', 'Кнопка избранного'),
    ('SendMessageButton', 'Кнопка отправки сообщения'),
    ('StarRating', 'Рейтинг звёздами (interactive)'),
    ('HotkeysProvider', 'Провайдер горячих клавиш'),
    ('InfiniteScroll', 'Бесконечная прокрутка'),
]

add_table(['Компонент', 'Описание'], components)

# ============================================================
# 25. БЕЗОПАСНОСТЬ
# ============================================================
add_h1('25. Безопасность')

add_h2('25.1 Авторизация')
add_bullet('JWT-токены с expiration (7 дней)')
add_bullet('Хэширование паролей bcryptjs (12 раундов)')
add_bullet('Двойное хранение: cookie + localStorage')
add_bullet('Middleware: JWT верификация через jose (Edge Runtime)')

add_h2('25.2 Middleware')
add_bullet('Проверка JWT токена (не только existence, но и validity)')
add_bullet('Невалидный токен → редирект + удаление cookie')
add_bullet('Проверка роли ADMIN для /admin маршрутов')
add_bullet('Совместимость с Edge Runtime (jose вместо jsonwebtoken)')

add_h2('25.3 Защита API')
add_bullet('Проверка токена в каждом защищённом API')
add_bullet('Проверка авторства (только автор может редактировать/удалять)')
add_bullet('Роль ADMIN для админ-эндпоинтов')

add_h2('25.4 Rate Limiting')
add_bullet('In-memory rate limiting (Map + cleanup каждую минуту)')
add_bullet('Login: 10 попыток/мин')
add_bullet('Register: 5 попыток/мин')
add_bullet('Upload: 20 файлов/мин')
add_bullet('Headers: X-RateLimit-Remaining, X-RateLimit-Reset')

add_h2('25.5 Валидация')
add_bullet('Zod-схемы для всех форм (register, login, post, comment, group, event, product, review, bookmark)')
add_bullet('Санитизация HTML input')
add_bullet('Проверка magic bytes при загрузке файлов')

add_h2('25.6 Загрузка файлов')
add_bullet('Ограничение размера: 10MB (изображения), 10MB (документы)')
add_bullet('Проверка MIME-типа')
add_bullet('Проверка magic bytes (JPEG, PNG, GIF, WebP)')
add_bullet('Блокировка опасных расширений (svg, php, exe, bat)')
add_bullet('Уникальные имена файлов (timestamp-random.ext)')
add_bullet('Хранение в public/uploads/')

add_h2('25.7 Git Safety')
add_bullet('.gitignore: .env, .env.local, prisma/dev.db, *.log, node_modules')
add_bullet('.env.example с плейсхолдерами (без реальных значений)')

# ============================================================
# 26. ТЕКУЩИЕ ПРОБЛЕМЫ И ПЛАНЫ
# ============================================================
add_h1('26. Текущие проблемы и планы')

add_h2('26.1 Исправлено (24.06.2026)')
add_bullet('JWT валидация в middleware через jose')
add_bullet('Проверка роли ADMIN для /admin маршрутов')
add_bullet('Удаление невалидного cookie')
add_bullet('Файловое хранилище public/uploads создано')

add_h2('26.2 Требует решения')
add_bullet('Credentials в .env — проверить что не отслеживается git')
add_bullet('Rate limiting in-memory — сбрасывается при рестарте')
add_bullet('Нет тестов')
add_bullet('Нет CI/CD')
add_bullet('Нет Docker')
add_bullet('Нет email верификации')
add_bullet('Нет сброса пароля')

add_h2('26.3 Планы улучшений')
add_bullet('Фаза 4: Тёмная тема, мобильная адаптация, email уведомления')
add_bullet('Фаза 5: Реалтайм уведомления (WebSocket/SSE)')
add_bullet('Фаза 6: Redis кэширование, CDN для изображений')
add_bullet('Фаза 7: Мобильное приложение (React Native)')
add_bullet('Фаза 8: Интеграция с 1С, CRM')
add_bullet('Фаза 9: API-платформа для мебельных компаний')

# ============================================================
# СОХРАНЕНИЕ
# ============================================================
output_path = r'E:\Код 2\Мебельный_портал_описание_возможностей.docx'
doc.save(output_path)
print(f'Документ сохранён: {output_path}')
