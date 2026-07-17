#!/usr/bin/env python3
"""Generate furniture enterprise documents: Act of Acceptance, Defects Act, Technical Specification."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import argparse


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def tune_styles(doc):
    body = doc.styles["Normal"]
    body.font.name = "Times New Roman"
    body.font.size = Pt(12)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.space_before = Pt(0)


def add_centered(doc, text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_body(doc, text, indent_first=True):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def add_sub(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def generate_acceptance_act(data, output_path):
    """Акт приёма-передачи мебели."""
    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    add_centered(doc, "АКТ", bold=True, size=16)
    add_centered(doc, "приёма-передачи мебельной продукции", bold=True, size=14)
    add_centered(doc, "(по договору №{} от {})".format(
        data.get("contract_number", "_____"),
        data.get("contract_date", "«___» _____________ 20___ г.")),
        bold=False, size=11)

    doc.add_paragraph()

    add_body(doc, "г. {}                                                                     "
                  "«___» _____________ 20___ г.".format(data.get("city", "_______________")), indent_first=False)

    doc.add_paragraph()

    add_body(doc,
        "{}, именуемое в дальнейшем «Исполнитель», в лице {} {}, "
        "действующего на основании {}, с одной стороны, и".format(
            data.get("company_full_name", "___________________________________________"),
            data.get("company_ceo_position", "Генерального директора"),
            data.get("company_ceo_name", "______________________________"),
            data.get("company_ceo_basis", "Устава"),
        ))

    add_body(doc,
        "{}, именуемое в дальнейшем «Заказчик», в лице {} {}, "
        "действующего на основании {}, с другой стороны, "
        "совместно именуемые «Стороны», а по отдельности — «Сторона», "
        "составили настоящий Акт о нижеследующем:".format(
            data.get("client_full_name", "___________________________________________"),
            data.get("client_ceo_position", "Генерального директора"),
            data.get("client_ceo_name", "______________________________"),
            data.get("client_ceo_basis", "Устава"),
        ))

    doc.add_paragraph()

    add_body(doc, "1. Исполнитель передал, а Заказчик принял следующую мебельную продукцию:")

    doc.add_paragraph()

    # Table with products
    table = doc.add_table(rows=1 + len(data.get("items", [])), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["№", "Наименование продукции", "Ед. изм.", "Кол-во", "Примечание"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Times New Roman"

    items = data.get("items", [
        {"name": "Кухонный гарнитур", "unit": "комплект", "qty": "1", "note": ""},
        {"name": "Шкаф-купе 3-створчатый", "unit": "шт.", "qty": "1", "note": ""},
        {"name": "Обеденный стол", "unit": "шт.", "qty": "1", "note": ""},
    ])

    for idx, item in enumerate(items, 1):
        row = table.rows[idx]
        values = [str(idx), item.get("name", ""), item.get("unit", ""), item.get("qty", ""), item.get("note", "")]
        for col_idx, val in enumerate(values):
            row.cells[col_idx].text = val
            for p in row.cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"

    doc.add_paragraph()

    add_body(doc, "2. Продукция передана в количестве и ассортименте, указанном в таблице выше.")

    add_body(doc, "3. Качество переданной продукции соответствует техническим требованиям и условиям Договора.")

    add_body(doc, "4. Комплектность поставки соответствует спецификации, указанной в Договоре.")

    add_body(doc, "5. Претензии по качеству, количеству и комплектности переданной продукции "
                  "могут быть предъявлены в течение 5 (пяти) рабочих дней с момента подписания настоящего Акта.")

    add_body(doc, "6. Настоящий Акт составлен в двух (двух) идентичных экземплярах, "
                  "по одному для каждой из Сторон, каждый из которых имеет равную юридическую силу.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_data = [
        ("ИСПОЛНИТЕЛЬ", "ЗАКАЗЧИК"),
        (data.get("company_full_name", ""), data.get("client_full_name", "")),
        ("Подпись: ___________", "Подпись: ___________"),
        ("Дата: «___» _____________ 20___ г.", "Дата: «___» _____________ 20___ г."),
    ]

    for row_idx, (left, right) in enumerate(sig_data):
        for col_idx, text in enumerate([left, right]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
        if row_idx == 0:
            for col_idx in range(2):
                for p in table.rows[row_idx].cells[col_idx].paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.save(output_path)
    print("Acceptance act saved to: {}".format(output_path))


def generate_defects_act(data, output_path):
    """Акт выявленных недостатков."""
    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    add_centered(doc, "АКТ", bold=True, size=16)
    add_centered(doc, "выявленных недостатков мебельной продукции", bold=True, size=14)
    add_centered(doc, "(по договору №{} от {})".format(
        data.get("contract_number", "_____"),
        data.get("contract_date", "«___» _____________ 20___ г.")),
        bold=False, size=11)

    doc.add_paragraph()

    add_body(doc, "г. {}                                                                     "
                  "«___» _____________ 20___ г.".format(data.get("city", "_______________")), indent_first=False)

    doc.add_paragraph()

    add_body(doc,
        "{}, именуемое в дальнейшем «Заказчик», в лице {} {}, "
        "действующего на основании {}, с одной стороны, и".format(
            data.get("client_full_name", "___________________________________________"),
            data.get("client_ceo_position", "Генерального директора"),
            data.get("client_ceo_name", "______________________________"),
            data.get("client_ceo_basis", "Устава"),
        ))

    add_body(doc,
        "{}, именуемое в дальнейшем «Исполнитель», в лице {} {}, "
        "действующего на основании {}, с другой стороны, "
        "составили настоящий Акт о выявленных недостатках:".format(
            data.get("company_full_name", "___________________________________________"),
            data.get("company_ceo_position", "Генерального директора"),
            data.get("company_ceo_name", "______________________________"),
            data.get("company_ceo_basis", "Устава"),
        ))

    doc.add_paragraph()

    add_body(doc, "1. В результате приёмки мебельной продукции по Договору №{} от {} "
                  "выявлены следующие недостатки:".format(
                      data.get("contract_number", "_____"),
                      data.get("contract_date", "«___» _____________ 20___ г."),
                  ))

    doc.add_paragraph()

    # Defects table
    defects = data.get("defects", [
        {"name": "Кухонный гарнитур", "defect": "Царапина на фасаде верхнего шкафа", "location": "Левая створка", "severity": "Существенный"},
        {"name": "Шкаф-купе", "defect": "Неработающий механизм направляющей", "location": "Правый ящик", "severity": "Существенный"},
        {"name": "Обеденный стол", "defect": "Несовпадение цвета столешницы", "location": "Столешница", "severity": "Незначительный"},
    ])

    table = doc.add_table(rows=1 + len(defects), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["№", "Наименование продукции", "Описание недостатка", "Место расположения", "Степень существенности"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = "Times New Roman"

    for idx, d in enumerate(defects, 1):
        row = table.rows[idx]
        values = [str(idx), d.get("name", ""), d.get("defect", ""), d.get("location", ""), d.get("severity", "")]
        for col_idx, val in enumerate(values):
            row.cells[col_idx].text = val
            for p in row.cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"

    doc.add_paragraph()

    add_body(doc, "2. Выявленные недостатки не позволяют использовать продукцию по назначению "
                  "в полном объёме / существенно影响уют качество использования.")

    add_body(doc, "3. Исполнитель обязан устранить выявленные недостатки за свой счёт "
                  "в течение {} рабочих дней с момента подписания настоящего Акта.".format(
                      data.get("fix_days", "14"),
                  ))

    add_body(doc, "4. В случае неустранения недостатков в указанные сроки Заказчик "
                  "вправе потребовать замены продукции или возврата уплаченных денежных средств.")

    add_body(doc, "5. Настоящий Акт составлен в двух (двух) идентичных экземплярах, "
                  "по одному для каждой из Сторон.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_data = [
        ("ЗАКАЗЧИК", "ИСПОЛНИТЕЛЬ"),
        (data.get("client_full_name", ""), data.get("company_full_name", "")),
        ("Подпись: ___________", "Подпись: ___________"),
        ("Дата: «___» _____________ 20___ г.", "Дата: «___» _____________ 20___ г."),
    ]

    for row_idx, (left, right) in enumerate(sig_data):
        for col_idx, text in enumerate([left, right]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
        if row_idx == 0:
            for col_idx in range(2):
                for p in table.rows[row_idx].cells[col_idx].paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.save(output_path)
    print("Defects act saved to: {}".format(output_path))


def generate_tz(data, output_path):
    """Техническое задание на изготовление мебели."""
    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    add_centered(doc, "ТЕХНИЧЕСКОЕ ЗАДАНИЕ", bold=True, size=16)
    add_centered(doc, "на изготовление мебельной продукции", bold=True, size=14)

    doc.add_paragraph()

    add_body(doc, "г. {}                                                                     "
                  "«___» _____________ 20___ г.".format(data.get("city", "_______________")), indent_first=False)

    doc.add_paragraph()

    # Section 1 - General info
    add_centered(doc, "1. ОБЩИЕ СВЕДЕНИЯ", bold=True, size=12)
    doc.add_paragraph()

    add_body(doc, "1.1. Заказчик: {}.".format(
        data.get("client_full_name", "___________________________________________")))

    add_body(doc, "1.2. Исполнитель: {}.".format(
        data.get("company_full_name", "___________________________________________")))

    add_body(doc, "1.3. Основание: Договор №{} от {}.".format(
        data.get("contract_number", "_____"),
        data.get("contract_date", "«___» _____________ 20___ г.")))

    add_body(doc, "1.4. Цель: изготовление мебельной продукции в соответствии с "
                  "настоящим Техническим заданием и условиями Договора.")

    doc.add_paragraph()

    # Section 2 - Requirements
    add_centered(doc, "2. ТРЕБОВАНИЯ К ПРОДУКЦИИ", bold=True, size=12)
    doc.add_paragraph()

    add_body(doc, "2.1. Мебельная продукция должна соответствовать следующим нормативным документам:")
    add_sub(doc, "- ГОСТ 19917-2014 «Мебельные изделия. Общие технические условия»")
    add_sub(doc, "- ГОСТ 20400-2017 «Продукция мебельного производства. Общие технические условия»")
    add_sub(doc, "- Технический регламент Таможенного союза ТР ТС 025/2012 «О безопасности мебельной продукции»")

    add_body(doc, "2.2. Материалы должны быть сертифицированы и соответствовать "
                  "требованиям экологической безопасности (класс эмиссии Е1 или выше).")

    add_body(doc, "2.3. Фурнитура должна быть от проверенных производителей (Blum, Hettich, Grass или аналоги).")

    add_body(doc, "2.4. Гарантийный срок эксплуатации — не менее 12 (двенадцати) месяцев.")

    doc.add_paragraph()

    # Section 3 - Product list
    add_centered(doc, "3. ПЕРЕЧЕНЬ ПРОДУКЦИИ", bold=True, size=12)
    doc.add_paragraph()

    items = data.get("items", [
        {"name": "Кухонный гарнитур", "dimensions": "длина 3000 мм, высота 2200 мм",
         "materials": "ЛДСП 18 мм, фасад МДФ эмаль, столешница искусственный камень",
         "qty": "1 комплект"},
        {"name": "Шкаф-купе 3-створчатый", "dimensions": "ширина 2400 мм, высота 2600 мм, глубина 600 мм",
         "materials": "ЛДСП 18 мм, зеркало 4 мм, направляющие Blum",
         "qty": "1 шт."},
        {"name": "Обеденный стол", "dimensions": "1600×900×750 мм",
         "materials": "Массив дуба, покрытие масло-воск",
         "qty": "1 шт."},
    ])

    table = doc.add_table(rows=1 + len(items), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["№", "Наименование", "Размеры (Д×Ш×В)", "Материалы", "Кол-во"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Times New Roman"

    for idx, item in enumerate(items, 1):
        row = table.rows[idx]
        values = [str(idx), item.get("name", ""), item.get("dimensions", ""),
                  item.get("materials", ""), item.get("qty", "")]
        for col_idx, val in enumerate(values):
            row.cells[col_idx].text = val
            for p in row.cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"

    doc.add_paragraph()

    # Section 4 - Additional requirements
    add_centered(doc, "4. ДОПОЛНИТЕЛЬНЫЕ ТРЕБОВАНИЯ", bold=True, size=12)
    doc.add_paragraph()

    add_body(doc, "4.1. Все изделия должны иметь защитную упаковку, предотвращающую повреждения "
                  "при транспортировке.")

    add_body(doc, "4.2. Каждое изделие должно быть промаркировано: наименование, "
                  "дата изготовления, номер заказа.")

    add_body(doc, "4.3. Исполнитель предоставляет фотоотчёт этапов производства по запросу Заказчика.")

    add_body(doc, "4.4. Доставка и монтаж осуществляются силами Исполнителя "
                  "в соответствии с условиями Договора.")

    doc.add_paragraph()

    # Section 5 - Timeline
    add_centered(doc, "5. СРОКИ ИЗГОТОВЛЕНИЯ", bold=True, size=12)
    doc.add_paragraph()

    add_body(doc, "5.1. Срок изготовления — {} календарных дней с момента подписания "
                  "настоящего Технического задания и поступления авансового платежа.".format(
                      data.get("production_days", "30"),
                  ))

    add_body(doc, "5.2. Заказчик вправе контролировать ход производства "
                  "в рабочие дни с 9:00 до 18:00.")

    doc.add_paragraph()

    # Signatures
    add_centered(doc, "6. ПОДПИСИ СТОРОН", bold=True, size=12)
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_data = [
        ("ЗАКАЗЧИК", "ИСПОЛНИТЕЛЬ"),
        (data.get("client_full_name", ""), data.get("company_full_name", "")),
        ("Подпись: ___________", "Подпись: ___________"),
        ("Дата: «___» _____________ 20___ г.", "Дата: «___» _____________ 20___ г."),
    ]

    for row_idx, (left, right) in enumerate(sig_data):
        for col_idx, text in enumerate([left, right]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
        if row_idx == 0:
            for col_idx in range(2):
                for p in table.rows[row_idx].cells[col_idx].paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.save(output_path)
    print("TZ saved to: {}".format(output_path))


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate furniture documents")
    parser.add_argument("--type", choices=["act", "defects", "tz"], required=True)
    parser.add_argument("--out", required=True, help="Output file path")
    args = parser.parse_args()

    data = {
        "city": "Москва",
        "contract_number": "001/2026",
        "contract_date": "«15» июля 2026 г.",
        "company_full_name": "ООО \"МебельПро\"",
        "company_ceo_position": "Генерального директора",
        "company_ceo_name": "Иванов Иван Иванович",
        "company_ceo_basis": "Устава",
        "client_full_name": "ООО \"Мебельный Портал\"",
        "client_ceo_position": "Генерального директора",
        "client_ceo_name": "Петров Пётр Петрович",
        "client_ceo_basis": "Устава",
        "items": [
            {"name": "Кухонный гарнитур", "unit": "комплект", "qty": "1", "note": "",
             "dimensions": "длина 3000 мм, высота 2200 мм",
             "materials": "ЛДСП 18 мм, фасад МДФ эмаль, столешница искусственный камень"},
            {"name": "Шкаф-купе 3-створчатый", "unit": "шт.", "qty": "1", "note": "",
             "dimensions": "ширина 2400 мм, высота 2600 мм, глубина 600 мм",
             "materials": "ЛДСП 18 мм, зеркало 4 мм, направляющие Blum"},
            {"name": "Обеденный стол", "unit": "шт.", "qty": "1", "note": "",
             "dimensions": "1600×900×750 мм",
             "materials": "Массив дуба, покрытие масло-воск"},
        ],
        "defects": [
            {"name": "Кухонный гарнитур", "defect": "Царапина на фасаде верхнего шкафа",
             "location": "Левая створка", "severity": "Существенный"},
            {"name": "Шкаф-купе", "defect": "Неработающий механизм направляющей",
             "location": "Правый ящик", "severity": "Существенный"},
            {"name": "Обеденный стол", "defect": "Несовпадение цвета столешницы",
             "location": "Столешница", "severity": "Незначительный"},
        ],
        "fix_days": "14",
        "production_days": "30",
    }

    if args.type == "act":
        generate_acceptance_act(data, args.out)
    elif args.type == "defects":
        generate_defects_act(data, args.out)
    elif args.type == "tz":
        generate_tz(data, args.out)
